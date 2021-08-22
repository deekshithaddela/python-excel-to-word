
from docx import Document
from docx.shared import Cm, Pt

word_document = Document()
document_name = 'new'

for article in [1,2,3]:
    # extracting text stats
    text_stats = {}
    text_stats['Article'] = 1
    text_stats['a'] = 1
    text_stats['b'] = 1
    
    # customizing the table
    table = word_document.add_table(0, 0) # we add rows iteratively
    table.style = 'TableGrid'
    w = [5,0.5,8]
    for i in w:
        table.add_column(Cm(i))
    
    for index, stat_item in enumerate(text_stats.items()):
        table.add_row()
        stat_name, stat_result = stat_item
        row = table.rows[index]
        row.cells[0].text = str(stat_name)
        row.cells[1].text = ':'
        row.cells[2].text = str(stat_result)
    p = word_document.add_paragraph()
    run = p.add_run()
    run.add_break()
word_document.save(document_name + '.docx')